import ollama
from IPython.display import display, Markdown
from docx import Document


def generate_itinerary(videos_df, duration, budget, travel_style):
    """
    Generates a detailed itinerary based on summarized content from videos.
    """
    # Prepare the input for the model
    summaries = videos_df['Summary'].tolist()
    num_days = duration  # Assuming duration is an integer representing the number of days
    
    # Create a prompt for generating the itinerary
    prompt = f"""
    As a travel guide expert, generate a detailed day-by-day itinerary for a traveler based on the following summaries:
    {summaries}
    
    The trip lasts for {num_days} days, and the travel style is {travel_style}. 
    Please include budget considerations of {budget} and make sure the activities are evenly distributed across the days.
    
    Format the itinerary with the day number and a list of activities for each day.
    """
    
    try:
        # Get response from the model
        response = ollama.chat(model="llama3.2", messages=[{'role':'user', 'content': prompt}])
        itinerary_md = response['message']['content']
        
    except Exception as e:
        return f"Error generating itinerary: {e}"
    
    return itinerary_md

def display_itinerary_with_markdown(videos_df, duration, budget, travel_style):
    """
    Displays the generated itinerary in a markdown format.
    """
    itinerary_text = generate_itinerary(videos_df, duration, budget, travel_style)
    markdown_output = f"### Your Travel Itinerary:\n\n{itinerary_text}"
    display(Markdown(markdown_output))

def save_itinerary_to_doc(itinerary, filename='itinerary.docx'):
    """
    Saves the generated itinerary to a DOC file.
    """
    doc = Document()
    doc.add_heading('Travel Itinerary', level=1)

    # Split the itinerary into lines and add to the document
    for line in itinerary.split('\n'):
        if line.strip():  # Avoid empty lines
            doc.add_paragraph(line.strip())

    doc.save(filename)
